
from docx import Document
import os

TEMPLATE_DIR = "templates"

def replace_in_doc(doc, replacements):
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)


# 1. FIX SALARY CERTIFICATE
sal_path = os.path.join(TEMPLATE_DIR, "Salary Certificate.docx")
if os.path.exists(sal_path):
    print(f"Fixing {sal_path}...")
    doc = Document(sal_path)
    
    replacements = {
        "Mr. Nischal Maharjan": "{{Name}}",  # Contains Title + Name
        "122938597": "{{PAN}}",
        "NP007": "{{EmployeeID}}",
        "Chief Operating Officer": "{{Role}}",
        "July 2023": "{{JoinMonthYear}}",
        "1,30,000.00": "{{Salary}}",
        "One lakh Thirty Thousand": "{{SalaryWords}}", # Case sensitive check
        "23-08-2024": "{{Date}}",
        " he ": " {{he}} ",
        " He ": " {{He}} "
    }
    
    # Special handling for case-sensitive SalaryWords if needed
    # The analyze output said: "One lakh Thirty Thousand\nonly" might be split?
    # Let's try matching unique substrings.
    
    doc_text = "\n".join([p.text for p in doc.paragraphs])
    
    # Apply replacements
    replace_in_doc(doc, replacements)
    
    # Save
    doc.save(sal_path)
    print("Salary Certificate fixed!")

# 2. FIX EXPERIENCE LETTER
exp_path = os.path.join(TEMPLATE_DIR, "Experience Letter.docx")
if os.path.exists(exp_path):
    print(f"Fixing {exp_path}...")
    doc = Document(exp_path)
    
    replacements_exp = {
        "Mr. Mukunda Sah": "{{Name}}",
        "Accounts and Admin Manager": "{{Role}}",
        "August 2023": "{{JoinMonthYear}}",
        "Jan 21, 2026": "{{Date}}", # Adjust if date format differs
        "January 21, 2026": "{{Date}}",
        " he ": " {{he}} ",
        " He ": " {{He}} ",
        " his ": " {{his}} ",
        " His ": " {{His}} ",
        "Maintaining complete financial records, including bank reconciliations and VAT return filings.": "{{Responsibilities}}",
        "Utilizing Tally and Xero for managing accou": "" 
    }
    
    replace_in_doc(doc, replacements_exp)
    
    # Also remove the other bullets if only one replaced.
    # Actually, simpler: finding the paragraph with "Maintaining..." and replacing it.
    for p in doc.paragraphs:
        if "Utilizing Tally" in p.text:
            p.text = "" # Clear it
        if "Maintaining complete" in p.text:
            p.text = "{{Responsibilities}}"
            
    doc.save(exp_path)
    print("Experience Letter fixed!")
