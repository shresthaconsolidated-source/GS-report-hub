import streamlit as st
import pandas as pd
import os
import io
import requests
from docx import Document
import re

st.set_page_config(page_title="HR Document Generator", page_icon="📝", layout="wide")

# --- AUTHENTICATION ---
if "hr_auth" not in st.session_state:
    st.session_state["hr_auth"] = False

if not st.session_state["hr_auth"]:
    st.title("🔒 HR Restricted Access")
    pwd = st.text_input("Enter Admin Password", type="password", key="hr_pwd")
    if st.button("Login"):
        if pwd == "Admin123":
            st.session_state["hr_auth"] = True
            st.rerun()
        else:
            st.error("Invalid Password")
    st.stop()


st.title("📝 HR Document Generator")
st.markdown("Generate official Profit / Experience letters linked to live Notion data.")

# API Configuration
API_URL = "https://hr-api.ashishoct34.workers.dev/api/hr"

# --- HELPER FUNCTIONS ---
def money_in_words(num):
    """Converts number to words (Indian Numbering System)"""
    try:
        num = int(float(str(num).replace(',', '')))
    except:
        return ""
    
    if num == 0: return "Zero"
    
    units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", 
             "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    
    def convert_upto_99(n):
        if n < 20: return units[n]
        return tens[n // 10] + (" " + units[n % 10] if n % 10 != 0 else "")
    
    def convert_upto_999(n):
        if n < 100: return convert_upto_99(n)
        return units[n // 100] + " Hundred" + (" and " + convert_upto_99(n % 100) if n % 100 != 0 else "")
    
    words = []
    
    # Crores
    if num >= 10000000:
        crores = num // 10000000
        words.append(convert_upto_99(crores) + " Crore")
        num %= 10000000
    
    # Lakhs
    if num >= 100000:
        lakhs = num // 100000
        words.append(convert_upto_99(lakhs) + " Lakh")
        num %= 100000

    # Thousands
    if num >= 1000:
        thousands = num // 1000
        words.append(convert_upto_99(thousands) + " Thousand")
        num %= 1000
        
    # Hundreds
    if num > 0:
        words.append(convert_upto_999(num))
        
    return " ".join(words) + " only"

def get_employees():
    """Fetch employees from Cloudflare Worker"""
    try:
        data = requests.get(API_URL).json()
        return data.get('employees', [])
    except Exception as e:
        st.error(f"Error fetching data: {e}")
        return []

def fill_template(template_path, data):
    """Replace placeholders in docx"""
    doc = Document(template_path)
    
    def replace_text(paragraph):
        for key, value in data.items():
            # Support {{Key}} and <Key> and [Key]
            patterns = [f"{{{{{key}}}}}", f"<{key}>", f"[{key}]"]
            for p in patterns:
                if p in paragraph.text:
                    paragraph.text = paragraph.text.replace(p, str(value))
    
    for p in doc.paragraphs:
        replace_text(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text(p)
                    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 1. LOAD DATA ---
with st.spinner("Fetching employee data from Notion..."):
    employees = get_employees()

if not employees:
    st.stop()

# --- 2. SELECT EMPLOYEE ---
employee_names = sorted([e.get('name', 'Unknown') for e in employees])
selected_name = st.selectbox("👤 Select Employee", employee_names)

# Get full employee object
employee = next((e for e in employees if e['name'] == selected_name), None)

if employee:
    # --- 3. SHOW DETAILS ---
    with st.expander("📄 View Employee Details", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.write(f"**Name:** {employee.get('name')}")
            st.write(f"**ID:** {employee.get('employee_id', 'N/A')}")
            st.write(f"**Designation:** {employee.get('designation', 'N/A')}")
        with col2:
            st.write(f"**Join Date:** {employee.get('joining_date', 'N/A')}")
            st.write(f"**Gender:** {employee.get('gender', 'N/A')}")
            st.write(f"**PAN:** {employee.get('pan', 'N/A')}")
        with col3:
             st.write(f"**Salary:** NPR {employee.get('last_salary', 0):,.2f}")
             st.write(f"**Email:** {employee.get('email', 'N/A')}")

    st.divider()

    # --- 4. DATA PREPARATION ---
    # Gender Logic
    gender = employee.get('gender', '').lower()
    is_female = "female" in gender
    
    title = "Miss" if is_female else "Mr."
    HeShe = "She" if is_female else "He"
    he_she = "she" if is_female else "he"
    HimHer = "Her" if is_female else "Him"
    him_her = "her" if is_female else "him"
    HisHer = "Her" if is_female else "His"
    his_her = "her" if is_female else "his"
    
    # Date Logic
    join_date_raw = employee.get('joining_date', '')
    join_month_year = join_date_raw
    try:
        # Try parsing ISO format or common formats
        if join_date_raw:
            ts = pd.to_datetime(join_date_raw)
            join_month_year = ts.strftime('%B %Y') # e.g. July 2023
            join_date_fmt = ts.strftime('%d-%m-%Y')
        else:
            join_date_fmt = ""
    except:
        join_date_fmt = join_date_raw
    
    # Current Date Format: January 29th, 2026
    def get_ordinal_date(date_obj):
        day = date_obj.day
        if 4 <= day <= 20 or 24 <= day <= 30:
            suffix = "th"
        else:
            suffix = ["st", "nd", "rd"][day % 10 - 1]
        return date_obj.strftime(f"%B {day}{suffix}, %Y")

    formatted_today = get_ordinal_date(pd.Timestamp.now())

    # Salary Logic
    salary_num = employee.get('last_salary', 0)
    salary_fmt = f"{salary_num:,.2f}"
    salary_words = money_in_words(salary_num)
    
    # 60/40 Breakdown
    basic_val = salary_num * 0.60
    allowance_val = salary_num * 0.40
    
    template_data = {
        "Name": f"{title} {employee.get('name', '')}", 
        "OnlyName": employee.get('name', ''),
        "Title": title,
        
        "he": he_she, "He": HeShe,
        "him": him_her, "Him": HimHer,
        "his": his_her, "His": HisHer,
        
        "EmployeeID": employee.get('employee_id', ''),
        "Role": employee.get('designation', ''),
        "Designation": employee.get('designation', ''),
        
        "JoinDate": join_date_fmt,          
        "JoinMonthYear": join_month_year,   
        
        "Salary": salary_fmt,
        "BasicSalary": f"{basic_val:,.2f}",
        "DearnessAllowance": f"{allowance_val:,.2f}",
        
        "SalaryWords": salary_words,
        "Department": employee.get('department', ''),
        "PAN": employee.get('pan', 'N/A'),
        "Date": formatted_today,
        "ReportingManager": employee.get('reporting_manager', 'Office Manager'),
        "Responsibilities": "" 
    }

    # Debug info for user
    # st.write(template_data) 
    
    with st.expander("ℹ️  Template Variables Guide"):
        st.write("""
        Use these placeholders in your Word document:
        - `{{Name}}` : Mr. John Doe
        - `{{Role}}` : Chief Operating Officer
        - `{{JoinDate}}` : 23-01-2026
        - `{{Salary}}` : 1,30,000.00
        - `{{ReportingManager}}` : Manager Name (or default)
        - `{{he}}` / `{{she}}` : Pronouns
        """)

    # --- 5. GENERATE DOCUMENTS (Notion Based) ---
    col_sal, col_exp = st.columns(2)
    
    TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")
    
    with col_sal:
        st.subheader("💰 Salary Certificate")
        st.info("Generates certificate using Notion data.")
        
        sal_template = "Salary Certificate.docx"
        sal_path = os.path.join(TEMPLATE_DIR, sal_template)
        
        if os.path.exists(sal_path):
            if st.button("Generate Salary Cert"):
                doc_buffer = fill_template(sal_path, template_data)
                st.download_button(
                    label="⬇️ Download DOCX",
                    data=doc_buffer,
                    file_name=f"Salary_Certificate_{employee['name']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error(f"Template '{sal_template}' not found.")

    with col_exp:
        st.subheader("📜 Experience Letter")
        st.info("Generates letter using Notion data.")
        
        responsibilities = st.text_area("Key Tasks (for Exp Letter)", height=150, placeholder="• Task 1\n• Task 2...")
        
        exp_data = template_data.copy()
        exp_data["Responsibilities"] = responsibilities
        
        exp_template = "Experience Letter.docx"
        exp_path = os.path.join(TEMPLATE_DIR, exp_template)
        
        if os.path.exists(exp_path):
             if st.button("Generate Exp Letter"):
                doc_buffer = fill_template(exp_path, exp_data)
                st.download_button(
                    label="⬇️ Download DOCX",
                    data=doc_buffer,
                    file_name=f"Experience_Letter_{employee['name']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error(f"Template '{exp_template}' not found.")

    st.divider()

# --- 6. MANUAL CONTRACT GENERATION (New Hires) ---
st.header("🤝 New Hire Contract (Manual Entry)")
st.caption("Use this for new employees who are not yet in Notion.")

mc_col1, mc_col2 = st.columns(2)

with mc_col1:
    m_name = st.text_input("Candidate Name", placeholder="Full Name")
    m_gender = st.selectbox("Gender", ["Male", "Female"])
    m_role = st.text_input("Position / Role", placeholder="e.g. Senior Counselor")
    m_department = st.text_input("Department", value="Administration") # New
    m_join_date = st.date_input("Joining Date")

with mc_col2:
    m_salary = st.number_input("Monthly Gross Salary (NPR)", min_value=0.0, step=1000.0)
    m_manager = st.text_input("Reporting Manager", value="Office Manager")
    m_signatory = st.text_input("Authorized Signatory", value="Managing Director") # New (Default to generic or allow name)
    m_responsibilities = st.text_area("Key Responsibilities", height=105, placeholder="• Task 1...")

if st.button("Generate New Hire Contract"):
    if not m_name:
        st.error("Please enter a name.")
    else:
        # 1. Gender Logic
        title = "Miss" if m_gender == "Female" else "Mr."
        he_she = "she" if m_gender == "Female" else "he"
        HeShe = "She" if m_gender == "Female" else "He"
        him_her = "her" if m_gender == "Female" else "him"
        HimHer = "Her" if m_gender == "Female" else "Him"
        his_her = "her" if m_gender == "Female" else "his"
        HisHer = "Her" if m_gender == "Female" else "His"
        
        # 2. Salary Logic (60/40)
        m_basic = m_salary * 0.60
        m_allowance = m_salary * 0.40
        m_salary_words = money_in_words(m_salary)
        
        # 3. Date Formatting
        formatted_join = get_ordinal_date(pd.Timestamp(m_join_date))
        
        m_data = {
            "Name": f"{title} {m_name}",
            "Title": title,
            "Role": m_role,
            "Department": m_department,
            "JoinDate": formatted_join,
            "Salary": f"{m_salary:,.2f}",
            "BasicSalary": f"{m_basic:,.2f}",
            "DearnessAllowance": f"{m_allowance:,.2f}",
            "SalaryWords": m_salary_words,
            "ReportingManager": m_manager,
            "AuthorizedSignatory": m_signatory,
            "Responsibilities": m_responsibilities,
            "Date": formatted_today,
            
            # Pronouns
            "he": he_she, "He": HeShe,
            "him": him_her, "Him": HimHer,
            "his": his_her, "His": HisHer
        }
        
        # 4. Generate
        cont_path = os.path.join(TEMPLATE_DIR, "Employment Contract Fixed.docx")
        if os.path.exists(cont_path):
             doc_buffer = fill_template(cont_path, m_data)
             st.success(f"Generated Contract for {m_name}!")
             st.download_button(
                label="⬇️ Download Contract",
                data=doc_buffer,
                file_name=f"Employment_Contract_{m_name.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
             )
        else:
            st.error("Contract template not found.")



