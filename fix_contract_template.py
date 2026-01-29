
from docx import Document
import os

TEMPLATE_PATH = "templates/Employment Contract Fixed - Copy.docx"

def fix_contract():
    if not os.path.exists(TEMPLATE_PATH):
        print("Template not found!")
        return

    doc = Document(TEMPLATE_PATH)
    print("Fixing Employment Contract (Copy version)...")

    # The Copy template is already mostly fixed with placeholders
    # We just need to clean up the double date issue: "Joining Date: {{Date}} {{JoinDate}}"
    
    for p in doc.paragraphs:
        # Fix the double placeholder on Joining Date line
        if "Joining Date: {{Date}} {{JoinDate}}" in p.text:
            p.text = "Joining Date: {{JoinDate}}"
        # Also catch any variations
        elif "Joining Date:" in p.text and "{{Date}}" in p.text and "{{JoinDate}}" in p.text:
            p.text = "Joining Date: {{JoinDate}}"

    # Also check tables for salary breakdown placeholders (should already be there)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Verify the salary table has the right placeholders
        try:
            if "Basic Salary" in table.cell(0,0).text and not "{{BasicSalary}}" in table.cell(0,1).text:
                table.cell(0,1).text = "NPR {{BasicSalary}}"
            
            if "Dearness" in table.cell(1,0).text and not "{{DearnessAllowance}}" in table.cell(1,1).text:
                table.cell(1,1).text = "NPR {{DearnessAllowance}}"
                
            if "Gross Salary" in table.cell(2,0).text and not "{{Salary}}" in table.cell(2,1).text:
                table.cell(2,1).text = "NPR {{Salary}}"
        except:
            pass  # Table might not have this structure

    NEW_PATH = "templates/Employment Contract Fixed.docx"
    doc.save(NEW_PATH)
    print(f"Employment Contract Fixed! Saved to {NEW_PATH}")

if __name__ == "__main__":
    fix_contract()
