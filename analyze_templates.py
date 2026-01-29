
import os
import re
from docx import Document

TEMPLATE_DIR = "templates"

def extract_placeholders(text):
    # Regex for common placeholder patterns
    patterns = [
        r"\{\{(.*?)\}\}",    # {{Name}}
        r"\[(.*?)\]",        # [Date]
        r"&lt;(.*?)&gt;",    # <Role> (if xml escaped)
        r"<(.*?)>"           # <Role>
    ]
    found = set()
    for p in patterns:
        matches = re.findall(p, text)
        for m in matches:
            found.add(m.strip())
    return found

def analyze_docx(filepath):
    try:
        doc = Document(filepath)
        all_text = []
        
        # Paragraphs
        for para in doc.paragraphs:
            all_text.append(para.text)
            
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
                    
        full_content = "\n".join(all_text)
        placeholders = extract_placeholders(full_content)
        
        print(f"\n--- FILE: {os.path.basename(filepath)} ---")
        if placeholders:
            print(f"Potential Placeholders Found: {placeholders}")
        else:
            print("No obvious {{}} or [] placeholders found. Here is a snippet of text:")
            print(full_content[:500] + "...")
            
    except Exception as e:
        print(f"Error reading {filepath}: {e}")

if __name__ == "__main__":
    if not os.path.exists(TEMPLATE_DIR):
        print("Templates dir not found")
    else:
        files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith('.docx')]
        for f in files:
            analyze_docx(os.path.join(TEMPLATE_DIR, f))
