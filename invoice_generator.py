import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        start={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        end={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def generate_invoice(supplier, invoice_date, invoice_no, period_str, amount, output_path):
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Recipient Created Tax Invoice")
    run.font.size = Pt(16)
    run.bold = True
    
    # Date and Invoice No
    doc.add_paragraph(f"Date: {invoice_date}")
    doc.add_paragraph(f"Invoice No: RCTI_SC_{invoice_no}")
    doc.add_paragraph("")
    
    # Supplier Details
    doc.add_paragraph(f"To Supplier : {supplier['name']}")
    doc.add_paragraph(f"Supplier ABN : {supplier['abn']}")
    doc.add_paragraph(f"Address : {supplier['address']}")
    doc.add_paragraph(f"Phone: {supplier['phone']}")
    doc.add_paragraph("")
    
    # Table for description and amount
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    # table.autofit = False
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(1.5)
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Amount (AUD)"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        
    # Content row
    row_cells = table.rows[1].cells
    p_desc = row_cells[0].paragraphs[0]
    p_desc.add_run("Subcontractor service to Global Select Education and Migration Pty Ltd\n")
    p_desc.add_run(f"({period_str})")
    
    amount_str = f"{float(amount):,.0f}" if float(amount).is_integer() else f"{float(amount):,.2f}"
    
    row_cells[1].text = amount_str
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Total row
    tot_cells = table.rows[2].cells
    tot_cells[0].text = "Total amount paid to your bank"
    run_tot = tot_cells[0].paragraphs[0].runs[0]
    run_tot.bold = True
    
    tot_cells[1].text = amount_str
    tot_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tot_cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph("")
    
    # Bank Details
    p_bank_header = doc.add_paragraph()
    r = p_bank_header.add_run("Amount paid to your Bank Account:")
    r.bold = True
    
    doc.add_paragraph(f"Bank Name: {supplier['bank_name']}")
    doc.add_paragraph(f"BSB No: {supplier['bsb']}")
    doc.add_paragraph(f"Account Number: {supplier['account_number']}")
    
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Test
    supplier = {
        "name": "Eve Miringching Magar",
        "abn": "37 836 121 282",
        "address": "2/27 Cypress Street Adelaide SA 4000",
        "phone": "0452 106 335",
        "bank_name": "Commonwealth Bank",
        "bsb": "065-000",
        "account_number": "1262 8218"
    }
    generate_invoice(supplier, "24 Oct 2025", "27", "6 Oct 2025 to 19 Oct 2025", 1560.0, "test.docx")
    print("Generated test.docx")
