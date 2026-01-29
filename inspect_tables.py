
from docx import Document
import os

TEMPLATE_PATH = "templates/Employment Contract Fixed.docx"

if os.path.exists(TEMPLATE_PATH):
    doc = Document(TEMPLATE_PATH)
    print(f"File found. Inspecting {len(doc.tables)} tables.")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            row_content = []
            for c_idx, cell in enumerate(row.cells):
                row_content.append(f"({r_idx},{c_idx}): '{cell.text.strip()}'")
            print(" | ".join(row_content))
else:
    print("File not found to inspect.")
